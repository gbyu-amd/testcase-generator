#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将 Word (.docx) 文档转换为 Markdown 文件，忽略所有图片。

转换策略：
- 标题 → Markdown 标题（#/##/### 等）
- 正文段落 → 普通文本段落
- 表格 → Markdown 表格
- 粗体/斜体 → Markdown 强调标记
- 列表（有序/无序）→ Markdown 列表
- 图片 → 完全忽略
- 输出文件默认写到 inputs/requirements/current_prd.md

依赖：python-docx（pip install python-docx）
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit(
        "缺少依赖：请先运行 pip install python-docx"
    )

from case_utils import (
    build_source_path,
    configure_output_encoding,
    ensure_under,
    project_root,
    write_text_file,
)


# Word 内置标题样式名（包含中文本地化名称）
HEADING_STYLE_PREFIXES = ("Heading", "heading", "标题")

# 列表样式关键字
LIST_BULLET_PREFIXES = ("List Bullet", "List bullet", "列表项目符号")
LIST_NUMBER_PREFIXES = ("List Number", "List number", "列表编号")


def _is_heading(paragraph) -> int:
    """返回标题级别（1-6），非标题返回 0。"""
    style_name = paragraph.style.name if paragraph.style else ""
    for prefix in HEADING_STYLE_PREFIXES:
        if style_name.startswith(prefix):
            suffix = style_name[len(prefix):].strip()
            if suffix.isdigit():
                return int(suffix)
            # "Heading" 不带数字 → 视为 H1
            if not suffix:
                return 1
    return 0


def _is_list_bullet(paragraph) -> bool:
    style_name = paragraph.style.name if paragraph.style else ""
    return any(style_name.startswith(p) for p in LIST_BULLET_PREFIXES)


def _is_list_number(paragraph) -> bool:
    style_name = paragraph.style.name if paragraph.style else ""
    return any(style_name.startswith(p) for p in LIST_NUMBER_PREFIXES)


def _has_inline_image(paragraph) -> bool:
    """检测段落中是否包含图片（drawing/pict）。"""
    xml = paragraph._element.xml
    return "w:drawing" in xml or "w:pict" in xml


def _paragraph_has_only_image(paragraph) -> bool:
    """段落只有图片，无其他文字内容。"""
    text = paragraph.text.strip()
    return not text and _has_inline_image(paragraph)


def _run_to_markdown(run) -> str:
    """将单个 run 的文字转为带格式的 Markdown 片段。"""
    text = run.text
    if not text:
        return ""
    # 转义 Markdown 特殊字符
    text = text.replace("\\", "\\\\")
    if run.bold and run.italic:
        return f"***{text}***"
    if run.bold:
        return f"**{text}**"
    if run.italic:
        return f"*{text}*"
    return text


def _paragraph_to_markdown(paragraph) -> str:
    """将段落文字（含 run 格式）转为 Markdown 字符串，忽略图片 run。"""
    parts: list[str] = []
    for run in paragraph.runs:
        # 跳过图片 run（含 drawing 或 pict 的 xml）
        run_xml = run._element.xml
        if "w:drawing" in run_xml or "w:pict" in run_xml:
            continue
        parts.append(_run_to_markdown(run))
    return "".join(parts).strip()


def _list_indent_level(paragraph) -> int:
    """估算列表缩进层级（0-based）。"""
    numpr = paragraph._element.find(qn("w:numPr"))
    if numpr is not None:
        ilvl = numpr.find(qn("w:ilvl"))
        if ilvl is not None:
            val = ilvl.get(qn("w:val"))
            if val and val.isdigit():
                return int(val)
    return 0


def _cell_to_text(cell) -> str:
    """将表格单元格的所有段落合并为单行文字。"""
    lines: list[str] = []
    for para in cell.paragraphs:
        if _paragraph_has_only_image(para):
            continue
        text = _paragraph_to_markdown(para)
        if text:
            lines.append(text)
    return " ".join(lines)


def _table_to_markdown(table) -> list[str]:
    """将 Word 表格转换为 Markdown 表格行列表。"""
    rows = table.rows
    if not rows:
        return []

    md_rows: list[list[str]] = []
    for row in rows:
        cells = [_cell_to_text(cell).replace("|", "\\|") for cell in row.cells]
        md_rows.append(cells)

    if not md_rows:
        return []

    col_count = max(len(r) for r in md_rows)
    # 补齐列数
    for row in md_rows:
        while len(row) < col_count:
            row.append("")

    lines: list[str] = []
    # 表头
    lines.append("| " + " | ".join(md_rows[0]) + " |")
    # 分隔行
    lines.append("| " + " | ".join(["---"] * col_count) + " |")
    # 数据行
    for row in md_rows[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return lines


def convert_docx_to_markdown(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    output_lines: list[str] = []
    number_tracker: dict[int, int] = {}
    prev_was_list = False

    for block in doc.element.body:
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag

        if tag == "p":
            # 找到对应的 paragraph 对象
            para = None
            for p in doc.paragraphs:
                if p._element is block:
                    para = p
                    break
            if para is None:
                continue

            # 只有图片的段落 → 跳过
            if _paragraph_has_only_image(para):
                prev_was_list = False
                continue

            heading_level = _is_heading(para)
            if heading_level:
                text = para.text.strip()
                if text:
                    if output_lines and output_lines[-1] != "":
                        output_lines.append("")
                    output_lines.append("#" * heading_level + " " + text)
                    output_lines.append("")
                number_tracker.clear()
                prev_was_list = False
                continue

            text = _paragraph_to_markdown(para)

            if _is_list_bullet(para):
                level = _list_indent_level(para)
                indent = "  " * level
                if text:
                    output_lines.append(f"{indent}- {text}")
                number_tracker.pop(level, None)
                prev_was_list = True
                continue

            if _is_list_number(para):
                level = _list_indent_level(para)
                indent = "  " * level
                number_tracker[level] = number_tracker.get(level, 0) + 1
                # 重置更深层级的计数
                for k in list(number_tracker.keys()):
                    if k > level:
                        del number_tracker[k]
                seq = number_tracker[level]
                if text:
                    output_lines.append(f"{indent}{seq}. {text}")
                prev_was_list = True
                continue

            # 普通段落
            if prev_was_list and text:
                output_lines.append("")
            if text:
                output_lines.append(text)
                output_lines.append("")
            prev_was_list = False

        elif tag == "tbl":
            # 找到对应的 table 对象
            table = None
            for t in doc.tables:
                if t._element is block:
                    table = t
                    break
            if table is None:
                continue

            if output_lines and output_lines[-1] != "":
                output_lines.append("")
            table_lines = _table_to_markdown(table)
            output_lines.extend(table_lines)
            output_lines.append("")
            prev_was_list = False

    # 去除末尾多余空行
    while output_lines and output_lines[-1] == "":
        output_lines.pop()

    return "\n".join(output_lines) + "\n"


def list_sections(markdown: str) -> list[tuple[int, str]]:
    """提取 Markdown 中所有标题，返回 (级别, 标题文字) 列表。"""
    sections: list[tuple[int, str]] = []
    for line in markdown.splitlines():
        m = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if m:
            sections.append((len(m.group(1)), m.group(2).strip()))
    return sections


def print_sections(sections: list[tuple[int, str]]) -> None:
    if not sections:
        print("未找到任何标题章节。")
        return
    print(f"共 {len(sections)} 个章节：\n")
    for level, title in sections:
        indent = "  " * (level - 1)
        print(f"{indent}{'#' * level} {title}")


def extract_section(markdown: str, section: str) -> str:
    """从完整 Markdown 中提取指定章节内容（模糊匹配标题）。"""
    lines = markdown.splitlines()
    start = None
    start_level = 0

    normalized_section = section.strip().lower()
    for i, line in enumerate(lines):
        m = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if m and normalized_section in m.group(2).strip().lower():
            start = i
            start_level = len(m.group(1))
            break

    if start is None:
        return ""

    end = len(lines)
    for i in range(start + 1, len(lines)):
        m = re.match(r"^(#{1,6})\s+", lines[i])
        if m and len(m.group(1)) <= start_level:
            end = i
            break

    return "\n".join(lines[start:end]).strip()


def parse_args() -> argparse.Namespace:
    root = project_root()
    default_output = root / "inputs" / "requirements" / "current_prd.md"
    parser = argparse.ArgumentParser(
        description="将 Word (.docx) 转换为 Markdown，忽略所有图片。"
    )
    parser.add_argument("docx", help="源 Word 文件路径，例如 inputs/requirements/raw_docs/<文件名>.docx")
    parser.add_argument(
        "--output",
        default=str(default_output),
        help=f"输出 Markdown 文件路径，默认 {default_output.relative_to(root)}",
    )
    parser.add_argument(
        "--overwrite",
        action="store_true",
        help="输出文件已存在时覆盖；不加此参数则询问确认",
    )
    parser.add_argument(
        "--section",
        default="",
        help="只提取指定章节（模糊匹配标题），配合 --print 使用",
    )
    parser.add_argument(
        "--print",
        action="store_true",
        help="将转换结果直接打印到标准输出，不写文件（供 AI 直接读取）",
    )
    parser.add_argument(
        "--list-sections",
        action="store_true",
        help="列出文档中所有章节标题，不做转换和输出",
    )
    return parser.parse_args()


def main() -> int:
    configure_output_encoding()
    args = parse_args()
    root = project_root()

    docx_path = Path(args.docx)
    if not docx_path.is_absolute():
        docx_path = root / docx_path
    docx_path = docx_path.resolve()

    if not docx_path.exists():
        print(f"错误：文件不存在：{docx_path}", file=sys.stderr)
        return 1
    if docx_path.suffix.lower() != ".docx":
        print(f"错误：仅支持 .docx 格式，收到：{docx_path.suffix}", file=sys.stderr)
        return 1

    content = convert_docx_to_markdown(docx_path)

    if args.list_sections:
        print_sections(list_sections(content))
        return 0

    if args.section:
        content = extract_section(content, args.section)
        if not content:
            print(f"错误：未找到章节「{args.section}」，请检查标题名称。", file=sys.stderr)
            return 1

    # --print 模式：直接输出内容，供 AI 读取，不写文件
    if args.print:
        print(content)
        return 0

    output_path = ensure_under(
        build_source_path(args.output, root), root, "输出文件"
    )

    if output_path.exists() and not args.overwrite:
        answer = input(f"输出文件已存在：{output_path}\n覆盖？(y/N) ").strip().lower()
        if answer != "y":
            print("已取消。")
            return 0

    output_path.parent.mkdir(parents=True, exist_ok=True)
    write_text_file(output_path, content, encoding="utf-8")

    line_count = content.count("\n")
    print(f"转换完成：")
    print(f"  源文件：{docx_path}")
    print(f"  输出：  {output_path}")
    print(f"  行数：  {line_count}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
